import os
import time
import shutil
import re
from pypdf import PdfReader
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from docx import Document

# Directorios de trabajo
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "leer_pdf")
OUTPUT_DIR = os.path.join(BASE_DIR, "resultados_word")
PROCESSED_DIR = os.path.join(BASE_DIR, "procesados")

def setup_directories():
    """Crea los directorios necesarios si no existen."""
    for folder in [INPUT_DIR, OUTPUT_DIR, PROCESSED_DIR]:
        if not os.path.exists(folder):
            os.makedirs(folder)
            print(f"[*] Directorio creado: {folder}")

def clean_text(text):
    """Limpia artefactos comunes de codificación de PDF (especialmente en catalán)."""
    replacements = [
        ('dÆ', "d'"), ('DÆ', "D'"), ('lÆ', "l'"), ('LÆ', "L'"),
        ('n║', 'nº'), ('‗', 'ò'), ('Ú', 'é'), ('Ë', 'Ó'), ('¾', 'ó'),
        ('Þ', 'é'), ('Ó', 'à'), ('└', 'À'), ('Ý', 'í'), ('þ', 'ç'),
        ('À', 'ò'), ('ò', '·'), ('¬', 'à')
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    
    # Limpiar excesos de espacios en blanco
    text = re.sub(r'[ \t]+', ' ', text)
    return text

def text_extraction(pdf_path):
    """Extrae el contenido de texto de un archivo PDF usando pypdf."""
    try:
        reader = PdfReader(pdf_path)
        text = []
        for page_num in range(len(reader.pages)):
            page_text = reader.pages[page_num].extract_text()
            if page_text:
                text.append(page_text)
        
        full_text = "\n".join(text)
        return clean_text(full_text)
    except Exception as e:
        print(f"[!] Error extrayendo texto de {pdf_path}: {e}")
        return ""

def analyze_pcap_criteria(text):
    """Analiza el PCAP buscando criterios de adjudicación estructurados."""
    results = {
        "Subjetivos (Juicio de Valor)": {},
        "Objetivos (Automáticos)": []
    }

    # 1. Extraer Sección Subjetivos
    # Buscamos el inicio de la sección 1) de forma muy tolerante
    sub_start_re = r"1\)\s*CRITERIS?\s+D.ADJUDICACI.\s+SOTMESOS\s+A\s+(?:UN\s+)?JUDICI\s+DE\s+VALOR"
    obj_start_re = r"2[.\s]*VALORACI.\s+DELS\s+CRITERIS\s+QUANTIFICABLES\s+AUTOM.TICAMENT"
    
    sub_match = re.search(sub_start_re, text, re.IGNORECASE)
    obj_match = re.search(obj_start_re, text, re.IGNORECASE)

    if sub_match and obj_match:
        sub_text = text[sub_match.end():obj_match.start()]
        
        # Segmentar por LOTS
        lot_blocks = re.split(r"(LOTS?\s+[\d,\s+i]+(?:\([^)]*\))?)", sub_text, flags=re.IGNORECASE)
        for i in range(1, len(lot_blocks), 2):
            lot_name = lot_blocks[i].strip()
            lot_content = lot_blocks[i+1].strip()
            
            # Extraer sub-criterios (A, B, C...)
            criteria = re.findall(r"([A-Z]\)\s*[^:]+:.*?Fins\s+\d+\s+punts\.)", lot_content, re.DOTALL | re.IGNORECASE)
            if not criteria:
                # Fallback: párrafos razonables que empiecen por A), B), C)...
                criteria = re.findall(r"(^[A-Z]\)\s*[A-Z\s]+:.*)", lot_content, re.MULTILINE | re.IGNORECASE)
            
            cleaned_criteria = [re.sub(r'\s+', ' ', c).strip() for c in criteria if len(c.strip()) > 10]
            results["Subjetivos (Juicio de Valor)"][lot_name] = cleaned_criteria

    # 2. Extraer Sección Objetivos
    if obj_match:
        # Buscamos hasta la siguiente sección principal o final
        end_match = re.search(r"CRITERIS\s+PER\s+A\s+LA\s+DETERMINACIÓ", text[obj_match.end():], re.IGNORECASE)
        if end_match:
            obj_text = text[obj_match.end():obj_match.end() + end_match.start()]
        else:
            obj_text = text[obj_match.end():obj_match.end() + 2000] # Buffer
            
        # Extraer líneas de precio y tiempo
        obj_criteria = re.findall(r"([A-Z]\)\s*[^.]+[\.]+\s*FINS\s+A\s+\d+\s+PUNTS)", obj_text, re.IGNORECASE)
        if not obj_criteria:
            obj_criteria = re.findall(r"([A-Z]\)\s*[^.]+\.?\s*FINS\s+A\s+\d+\s+PUNTS)", obj_text, re.IGNORECASE)
        
        results["Objetivos (Automáticos)"] = [re.sub(r'\s+', ' ', c.strip()) for c in obj_criteria]

    if not results["Subjetivos (Juicio de Valor)"] and not results["Objetivos (Automáticos)"]:
        return {"error": "No se detectó la estructura de criterios esperada. Revise el PDF."}

    return results

def create_word_report(filename, analysis_results):
    """Genera un informe Word estructurado con los criterios encontrados."""
    doc = Document()
    doc.add_heading(f"Criterios de Adjudicación: {filename}", 0)

    if "error" in analysis_results:
        doc.add_paragraph(analysis_results["error"])
    else:
        # Sección Subjetivos
        doc.add_heading("1. Criterios Subjectivos (Juicio de Valor)", level=1)
        subjetivos = analysis_results.get("Subjetivos (Juicio de Valor)", {})
        if not subjetivos:
            doc.add_paragraph("No se detectaron criterios subjetivos.")
        else:
            for lot, criteria in subjetivos.items():
                doc.add_heading(lot, level=2)
                if criteria:
                    for c in criteria:
                        doc.add_paragraph(c, style='List Bullet')
                else:
                    doc.add_paragraph("Sin especificaciones detalladas detectadas para este lote.")

        # Sección Objetivos
        doc.add_heading("2. Criterios Objetivos (Automáticos)", level=1)
        objetivos = analysis_results.get("Objetivos (Automáticos)", [])
        if not objetivos:
            doc.add_paragraph("No se detectaron criterios objetivos.")
        else:
            for opt in objetivos:
                doc.add_paragraph(opt, style='List Bullet')

    output_path = os.path.join(OUTPUT_DIR, f"{filename}.docx")
    doc.save(output_path)
    print(f"[+] Informe generado exitosamente: {output_path}")

def process_single_pdf(pdf_path):
    """Lógica integral para procesar un documento a su llegada."""
    filename = os.path.basename(pdf_path)
    name_only, _ = os.path.splitext(filename)
    
    print(f"\n[>] Procesando nuevo PDF: {filename}")
    
    # Pequeña espera para evitar leer un archivo a medio copiar (Race condition)
    time.sleep(2) 

    text = text_extraction(pdf_path)
    if not text:
        print(f"[!] El archivo no contiene texto legible: {filename}")
        return

    print("[*] Analizando criterios de adjudicación...")
    analysis = analyze_pcap_criteria(text)
    
    print("[*] Generando documento Word...")
    create_word_report(name_only, analysis)

    # Mover el pdf a la carpeta de procesados para limpiar la entrada
    try:
        dest_path = os.path.join(PROCESSED_DIR, filename)
        shutil.move(pdf_path, dest_path)
        print(f"[*] Archivo movido a procesados: {filename}")
    except Exception as e:
        print(f"[!] Hubo un error moviendo el archivo: {e}")

class PCAPHandler(FileSystemEventHandler):
    """Escucha la creación de archivos nuevos en el directorio mapeado."""
    def on_created(self, event):
        if not event.is_directory and event.src_path.lower().endswith(".pdf"):
            process_single_pdf(event.src_path)

def start_watching():
    """Incia el loop de Watchdog sobre el directorio INPUT_DIR."""
    setup_directories()
    
    # Comprobar si hay archivos residuales preexistentes al iniciar
    for file in os.listdir(INPUT_DIR):
        if file.lower().endswith(".pdf"):
            process_single_pdf(os.path.join(INPUT_DIR, file))

    event_handler = PCAPHandler()
    observer = Observer()
    observer.schedule(event_handler, INPUT_DIR, recursive=False)
    observer.start()
    
    print(f"\n[*] Modo 'Ingeniero Senior' activo.")
    print(f"[*] Carpeta de escucha activa: {INPUT_DIR}")
    print("[*] Esperando que se suelten archivos PDF... (Aplica Ctrl+C para detener)")
    
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
        print("\n[!] Deteniendo el proceso...")
    
    observer.join()

if __name__ == '__main__':
    start_watching()
