import streamlit as st
from pypdf import PdfReader
from docx import Document
import io
import re

st.set_page_config(page_title="Extractor PCAP", page_icon="📄", layout="centered")

def text_extraction(pdf_file):
    """Extrae texto de un archivo PDF en memoria."""
    try:
        reader = PdfReader(pdf_file)
        text_lines = []
        for page_num in range(len(reader.pages)):
            page_text = reader.pages[page_num].extract_text()
            if page_text:
                text_lines.extend(page_text.split('\n'))
        return text_lines
    except Exception as e:
        return []

def analyze_pcap_criteria_lines(lines):
    """
    Analiza heurísticamente el PCAP línea por línea buscando la sección de criterios.
    Soporte bilingüe: Español (ES) y Catalán (CAT).
    """
    subjetivos = []
    objetivos = []
    
    # Flags para detectar si estamos dentro de un bloque interesante
    in_criteria_section = False
    buffer = []
    
    keyword_criteria = ["criterios de adjudicación", "criteris d'adjudicació", "criteris d’adjudicació"]
    keyword_subj = ["juicio de valor", "judici de valor", "subjetivo", "subjectiu", "no automátic", "no automàtic"]
    keyword_obj = ["fórmula", "fórmules", "automático", "automàtic", "objetivo", "objectiu"]
    
    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        
        # Detectar el inicio de la sección de criterios
        if any(kw in line_lower for kw in keyword_criteria):
            in_criteria_section = True
            
        if in_criteria_section:
            buffer.append(line.strip())
            
            # Si hemos acumulado bastantes lineas, busquemos los criterios
            if len(buffer) > 50: 
                # Solo procesar un pequeño bloque alrededor si no encontramos un "fin" de sección claro
                pass
                
        # Clasificar línea por línea si hay palabras clave
        # Esto es más robusto cuando el texto extraído pierde formato
        if any(kw in line_lower for kw in keyword_subj) and len(line_lower) > 20:
            # Añadir contexto: la línea anterior y posterior
            ctx = " ".join([l.strip() for l in lines[max(0, i-1):min(len(lines), i+2)]])
            # Evitar duplicados
            if ctx not in subjetivos:
                subjetivos.append(ctx)

        if any(kw in line_lower for kw in keyword_obj) and len(line_lower) > 20:
            ctx = " ".join([l.strip() for l in lines[max(0, i-1):min(len(lines), i+2)]])
            if ctx not in objetivos:
                objetivos.append(ctx)

    return {
        "Subjetivos / No Automáticos (Juicios de Valor)": subjetivos[:15], 
        "Objetivos / Automáticos (Fórmulas)": objetivos[:15]
    }

def create_word_buffer(filename, analysis_results):
    """Genera un .docx en memoria y devuelve el buffer para descargar."""
    doc = Document()
    doc.add_heading(f"Análisis PCAP: {filename}", 0)

    for title, criteria in analysis_results.items():
        doc.add_heading(title, level=1)
        if criteria:
            for c in criteria:
                doc.add_paragraph(c, style='List Bullet')
        else:
            doc.add_paragraph("No se detectaron criterios explícitos en esta categoría.")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- INTERFAZ STREAMLIT ---
st.title("📄 Extractor de Criterios PCAP (Word)")
st.markdown("""
<style>
.stApp {
    background-color: #f8f9fa;
}
</style>
""", unsafe_allow_html=True)

st.write("Sube tu Pliego (PDF) y extraeremos automáticamente los criterios de adjudicación (Soporta Castellano y Catalán).")

uploaded_file = st.file_uploader("Elige un archivo PDF", type="pdf")

if uploaded_file is not None:
    filename = uploaded_file.name
    st.info(f"Procesando: {filename}...")
    
    with st.spinner("Extrayendo texto y analizando criterios..."):
        lines = text_extraction(uploaded_file)
        
        if not lines:
            st.error("No se pudo extraer texto. Si es un PDF escaneado (imagen), requiere módulo OCR avanzado.")
        else:
            analysis = analyze_pcap_criteria_lines(lines)
            
            # Mostrar preview rápida
            for category, items in analysis.items():
                with st.expander(f"👁️ Vista Previa: {category}"):
                    if items:
                        for item in items:
                            st.write(f"- {item}")
                    else:
                        st.write("No encontrados.")
            
            # Generar Word en memoria
            word_buffer = create_word_buffer(filename, analysis)
            
            st.success("¡Informe generado con éxito!")
            
            # Botón de descarga
            st.download_button(
                label="📥 Descargar Informe en Word (.docx)",
                data=word_buffer,
                file_name=f"{filename.replace('.pdf', '')}_Criterios.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
