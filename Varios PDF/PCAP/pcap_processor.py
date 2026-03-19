import os
import time
import shutil
import re
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Directorios de trabajo
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, "leer_pdf")
OUTPUT_DIR = os.path.join(BASE_DIR, "resultados_word")
PROCESSED_DIR = os.path.join(BASE_DIR, "procesados")


def setup_directories():
    for folder in [INPUT_DIR, OUTPUT_DIR, PROCESSED_DIR]:
        if not os.path.exists(folder):
            os.makedirs(folder)


def clean_text(text):
    """Limpia artefactos de codificación comunes en PDFs catalanes/castellanos."""
    replacements = [
        ('dÆ', "d'"), ('DÆ', "D'"), ('lÆ', "l'"), ('LÆ', "L'"),
        ('nÆ', "n'"), ('sÆ', "s'"), ('NÆ', "N'"), ('SÆ', "S'"),
        ('n║', 'nº'), ('‗', 'ò'), ('Ú', 'é'), ('Ë', 'Ó'), ('¾', 'ó'),
        ('Þ', 'é'), ('Ý', 'í'), ('þ', 'ç'), ('¬', 'à'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    text = re.sub(r'[ \t]+', ' ', text)
    # Corregir apóstrofes seguidos de minúscula donde debería ser mayúscula (D'eS → D'ÚS)
    text = re.sub(r"D'eS", "D'ÚS", text)
    text = re.sub(r"d'eS", "d'ÚS", text)
    text = re.sub(r"d'es\b", "d'ús", text)
    return text


def extract_text(pdf_source):
    """Extrae texto de un PDF (path o file-like). Devuelve texto limpio completo."""
    try:
        reader = PdfReader(pdf_source)
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        return clean_text("\n".join(pages))
    except Exception as e:
        print(f"[!] Error extrayendo texto: {e}")
        return ""


# ---------------------------------------------------------------------------
#  PARSER DE CRITERIOS
# ---------------------------------------------------------------------------

# Patrones bilingües para secciones
_RE_SUBJ_START = re.compile(
    r'(?:CRITERIS\s+D.ADJUDICACI.\s+SOTMESOS\s+A\s+(?:UN\s+)?JUDICI\s+DE\s+VALOR'
    r'|CRITERIOS\s+DE\s+ADJUDICACI.N\s+SOMETIDOS\s+A\s+(?:UN\s+)?JUICIO\s+DE\s+VALOR)',
    re.IGNORECASE
)

_RE_OBJ_START = re.compile(
    r'(?:VALORACI.\s+DELS\s+CRITERIS\s+QUANTIFICABLES\s+AUTOM.TICAMENT'
    r'|CRITERIS\s+AVALUABLES\s+DE\s+FORMA\s+AUTOM.TICA'
    r'|CRITERIOS\s+EVALUABLES\s+DE\s+FORMA\s+AUTOM.TICA'
    r'|VALORACI.N\s+DE\s+(?:LOS\s+)?CRITERIOS\s+CUANTIFICABLES\s+AUTOM.TICAMENTE)',
    re.IGNORECASE
)

_RE_END = re.compile(
    r'(?:CRITERIS\s+PER\s+A\s+LA\s+DETERMINACI.'
    r'|CRITERIOS\s+PARA\s+LA\s+DETERMINACI.)',
    re.IGNORECASE
)

_RE_LOT_HEADER = re.compile(
    r'^(LOTS?\s+[\d,\s]+(?:i\s+\d+)?)\s*\(([^)]+)\)',
    re.IGNORECASE | re.MULTILINE
)

_RE_CRITERION = re.compile(
    r'([A-Z])\)\s*(.+?)(?:\.?\s*[Ff]ins\s+(?:a\s+)?(\d+)\s+punts'
    r'|\.?\s*[Hh]asta\s+(\d+)\s+puntos)',
    re.DOTALL
)

_RE_POINTS_INLINE = re.compile(
    r'[Ff]ins\s+(?:a\s+)?(\d+)\s+punts|[Hh]asta\s+(\d+)\s+puntos',
    re.IGNORECASE
)

_RE_RANGE_LINE = re.compile(
    r'^-\s*(Alta|Moderada|Baixa|Moderat|Baix|Baja|Alto|Bajo)\s*:\s*(.+)',
    re.IGNORECASE | re.MULTILINE
)


def _find_section(text, start_re, end_re):
    """Devuelve el bloque de texto entre start_re y end_re."""
    m_start = start_re.search(text)
    if not m_start:
        return None, -1, -1
    # Buscar el final después del inicio
    m_end = end_re.search(text, m_start.end())
    end_pos = m_end.start() if m_end else len(text)
    return text[m_start.start():end_pos], m_start.start(), end_pos


def _parse_subjective_section(text):
    """Parsea la sección de criterios subjetivos, devolviendo estructura por lotes."""
    # Buscar el bloque "1) CRITERIS..." que marca el inicio real de lotes
    m_numbered = re.search(
        r'1\)\s*CRITERIS\s+D.ADJUDICACI.\s+SOTMESOS|1\)\s*CRITERIOS\s+DE\s+ADJUDICACI.N\s+SOMETIDOS',
        text, re.IGNORECASE
    )
    if m_numbered:
        text = text[m_numbered.end():]

    # Extraer puntuación máxima total
    m_total = re.search(r'FINS\s+A\s+(\d+)\s+PUNTS|HASTA\s+(\d+)\s+PUNTOS', text, re.IGNORECASE)
    max_subj = int(m_total.group(1) or m_total.group(2)) if m_total else None

    # Encontrar todos los headers de lote
    lot_matches = list(_RE_LOT_HEADER.finditer(text))
    if not lot_matches:
        return {'max_points': max_subj, 'lots': [], 'raw': text.strip()}

    lots = []
    for i, lm in enumerate(lot_matches):
        lot_id = lm.group(1).strip()
        lot_desc = lm.group(2).strip()
        start = lm.end()
        end = lot_matches[i + 1].start() if i + 1 < len(lot_matches) else len(text)
        lot_text = text[start:end]

        # Buscar criterios A), B), C)... dentro del bloque del lote
        criteria = _parse_criteria_block(lot_text)
        lots.append({
            'id': lot_id,
            'description': lot_desc,
            'criteria': criteria
        })

    return {'max_points': max_subj, 'lots': lots}


def _parse_criteria_block(block_text):
    """Extrae criterios individuales (A, B, C...) de un bloque de texto."""
    # Dividir por letras de criterio: A), B), C)...
    parts = re.split(r'\n\s*([A-Z])\)\s+', block_text)
    criteria = []

    # parts[0] es texto antes del primer criterio, luego pares (letra, contenido)
    for j in range(1, len(parts) - 1, 2):
        letter = parts[j]
        content = parts[j + 1].strip()

        # Extraer nombre y puntos
        # El nombre es lo que hay antes de los dos puntos o del "Fins/Hasta"
        name_match = re.match(r'([^:]+?)(?:\s*:\s*|\s*\.\s*)', content)
        name = name_match.group(1).strip() if name_match else content[:60]

        # Extraer puntuación máxima
        pts_match = _RE_POINTS_INLINE.search(content)
        max_pts = int(pts_match.group(1) or pts_match.group(2)) if pts_match else None

        # Extraer descripción (primera frase antes de los rangos)
        desc_end = content.find('\n-')
        if desc_end == -1:
            desc_end = len(content)
        description = re.sub(r'\s+', ' ', content[:desc_end]).strip()

        # Extraer rangos (Alta/Moderada/Baixa) — pueden ser multilínea
        ranges = []
        # Dividir el contenido en bloques por "- " al inicio de línea
        range_blocks = re.split(r'\n\s*-\s+', content)
        for rb in range_blocks[1:]:  # Saltar el primer bloque (descripción)
            rb_clean = re.sub(r'\s+', ' ', rb).strip()
            # Detectar si es un rango con nivel
            level_match = re.match(
                r'(Alta|Moderada|Baixa|Moderat|Baix|Baja|Alto|Bajo)\s*:\s*(.*)',
                rb_clean, re.IGNORECASE
            )
            if level_match:
                level = level_match.group(1).strip()
                detail = level_match.group(2).strip()
            else:
                # Rango sin nivel explícito (e.g., presentación con sub-bullets)
                level = ""
                detail = rb_clean
            # Extraer rango de puntos del texto completo del bloque
            pts_range = re.search(
                r'[Dd]e\s+([\d.,]+)\s+a\s+([\d.,]+)\s+punts'
                r'|[Dd]e\s+([\d.,]+)\s+a\s+([\d.,]+)\s+puntos'
                r'|[Ff]ins\s+a\s+(\d+)\s+punts'
                r'|[Hh]asta\s+(\d+)\s+puntos'
                r'|[Dd]e\s+(\d+)\s+a\s+([\d.,]+)\s+punts',
                detail
            )
            ranges.append({
                'level': level,
                'detail': detail,
                'raw_points': pts_range.group(0) if pts_range else ''
            })

        criteria.append({
            'letter': letter,
            'name': name,
            'max_points': max_pts,
            'description': description,
            'ranges': ranges
        })

    return criteria


def _parse_objective_section(text):
    """Parsea la sección de criterios objetivos/automáticos."""
    m_total = re.search(r'M.XIM\s+(\d+)\s+PUNTS|M.XIMO?\s+(\d+)\s+PUNTOS', text, re.IGNORECASE)
    max_obj = int(m_total.group(1) or m_total.group(2)) if m_total else None

    criteria = []

    # Buscar criterios A), B)...
    parts = re.split(r'\n\s*([A-Z])\)\.?\s+', text)
    for j in range(1, len(parts) - 1, 2):
        letter = parts[j]
        content = parts[j + 1].strip()

        # Nombre y puntos
        first_line = content.split('\n')[0]
        pts_match = _RE_POINTS_INLINE.search(first_line)
        max_pts = int(pts_match.group(1) or pts_match.group(2)) if pts_match else None
        name = re.sub(r'\.?\s*[Ff]ins\s+.*', '', first_line).strip()
        name = re.sub(r'\.?\s*[Hh]asta\s+.*', '', name).strip()

        # Cuerpo completo
        body = re.sub(r'\s+', ' ', content).strip()

        # Extraer tiers de puntuación si existen (e.g., "2 dies = 5 punts")
        tiers = re.findall(
            r'-\s*(.+?=\s*\d+\s*punts?\.?)',
            content, re.IGNORECASE
        )
        if not tiers:
            tiers = re.findall(
                r'-\s*(.+?=\s*\d+\s*puntos?\.?)',
                content, re.IGNORECASE
            )

        criteria.append({
            'letter': letter,
            'name': name,
            'max_points': max_pts,
            'body': body,
            'tiers': tiers
        })

    return {'max_points': max_obj, 'criteria': criteria}


def analyze_pcap(text):
    """Análisis completo del PCAP. Devuelve estructura de criterios."""
    results = {
        'subjective': None,
        'objective': None,
        'warnings': []
    }

    # Encontrar sección subjetiva
    subj_text, s_start, s_end = _find_section(text, _RE_SUBJ_START, _RE_OBJ_START)
    if subj_text:
        results['subjective'] = _parse_subjective_section(subj_text)
    else:
        results['warnings'].append("No se ha encontrado la sección de criterios subjetivos (Judici de Valor / Juicio de Valor).")

    # Encontrar sección objetiva
    obj_text, o_start, o_end = _find_section(text, _RE_OBJ_START, _RE_END)
    if obj_text:
        results['objective'] = _parse_objective_section(obj_text)
    else:
        results['warnings'].append("No se ha encontrado la sección de criterios objetivos (Automátics / Automáticos).")

    if not results['subjective'] and not results['objective']:
        results['warnings'].append("No se ha detectado ninguna estructura de criterios de adjudicación en el documento.")

    return results


# ---------------------------------------------------------------------------
#  GENERADOR DE WORD
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def _add_styled_table(doc, headers, rows, header_color="004A99"):
    """Crea una tabla con estilo corporativo."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Cabecera
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    # Filas de datos
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                _set_cell_shading(cell, "F2F2F2")

    return table


def create_word_report(filename, analysis, output_path=None):
    """Genera un informe Word estructurado y limpio."""
    doc = Document()

    # --- Estilos globales ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # --- TÍTULO ---
    h = doc.add_heading(f"Criteris d'Adjudicació", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(filename)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 74, 153)

    subj = analysis.get('subjective')
    obj = analysis.get('objective')
    warnings = analysis.get('warnings', [])

    # --- RESUMEN ---
    doc.add_heading("Resum / Resumen", level=1)
    summary_rows = []
    if subj and subj.get('max_points'):
        summary_rows.append(["Criteris subjectius (Judici de Valor)", f"Fins a {subj['max_points']} punts"])
    if obj and obj.get('max_points'):
        summary_rows.append(["Criteris objectius (Automàtics)", f"Fins a {obj['max_points']} punts"])
    total = (subj.get('max_points', 0) or 0) + (obj.get('max_points', 0) or 0)
    if total:
        summary_rows.append(["TOTAL", f"{total} punts"])
    if summary_rows:
        _add_styled_table(doc, ["Tipus", "Puntuació màxima"], summary_rows)
    doc.add_paragraph("")

    # --- AVISOS ---
    if warnings:
        doc.add_heading("Avisos", level=1)
        for w in warnings:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠ {w}")
            run.font.color.rgb = RGBColor(200, 50, 0)

    # --- SECCIÓN SUBJETIVA ---
    if subj:
        pts_text = f" (Fins a {subj['max_points']} punts)" if subj.get('max_points') else ""
        doc.add_heading(f"1. Criteris Subjectius — Judici de Valor{pts_text}", level=1)

        lots = subj.get('lots', [])
        if lots:
            for lot in lots:
                doc.add_heading(f"{lot['id']} ({lot['description']})", level=2)

                if lot['criteria']:
                    # Tabla resumen del lote
                    rows = []
                    for c in lot['criteria']:
                        pts_str = f"{c['max_points']} punts" if c['max_points'] else "—"
                        rows.append([f"{c['letter']})", c['name'], pts_str])
                    _add_styled_table(doc, ["", "Criteri", "Punts màx."], rows, header_color="2E75B6")
                    doc.add_paragraph("")

                    # Detalle de cada criterio con rangos
                    for c in lot['criteria']:
                        pts_str = f" — Fins {c['max_points']} punts" if c['max_points'] else ""
                        p = doc.add_paragraph()
                        run = p.add_run(f"{c['letter']}) {c['name']}{pts_str}")
                        run.font.bold = True
                        run.font.size = Pt(10)

                        if c['ranges']:
                            for rng in c['ranges']:
                                p = doc.add_paragraph(style='List Bullet')
                                run_level = p.add_run(f"{rng['level']}: ")
                                run_level.font.bold = True
                                run_level.font.size = Pt(9)
                                run_detail = p.add_run(rng['detail'])
                                run_detail.font.size = Pt(9)
                else:
                    doc.add_paragraph("No s'han detectat criteris detallats per a aquest lot.")
        elif subj.get('raw'):
            doc.add_paragraph(subj['raw'])

    # --- SECCIÓN OBJETIVA ---
    if obj:
        pts_text = f" (Fins a {obj['max_points']} punts)" if obj.get('max_points') else ""
        doc.add_heading(f"2. Criteris Objectius — Automàtics{pts_text}", level=1)

        for c in obj.get('criteria', []):
            pts_str = f" — Fins {c['max_points']} punts" if c['max_points'] else ""
            doc.add_heading(f"{c['letter']}) {c['name']}{pts_str}", level=2)

            # Cuerpo descriptivo (sin los tiers)
            body_clean = c.get('body', '')
            # Quitar los tiers del body para no duplicar
            for tier in c.get('tiers', []):
                body_clean = body_clean.replace(tier, '')
            body_clean = re.sub(r'-\s+', '', body_clean).strip()
            body_clean = re.sub(r'\s+', ' ', body_clean)
            if body_clean:
                doc.add_paragraph(body_clean)

            # Tiers como tabla si existen
            if c.get('tiers'):
                tier_rows = []
                for tier_text in c['tiers']:
                    parts = tier_text.split('=')
                    if len(parts) == 2:
                        tier_rows.append([parts[0].strip(), parts[1].strip()])
                    else:
                        tier_rows.append([tier_text, ""])
                if tier_rows:
                    _add_styled_table(doc, ["Condició", "Puntuació"], tier_rows, header_color="2E75B6")
                    doc.add_paragraph("")

    # --- Guardar ---
    if output_path:
        doc.save(output_path)
        return output_path
    else:
        import io
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio


# ---------------------------------------------------------------------------
#  WATCHDOG (modo CLI)
# ---------------------------------------------------------------------------

def process_single_pdf(pdf_path):
    filename = os.path.basename(pdf_path)
    name_only, _ = os.path.splitext(filename)

    print(f"\n[>] Procesando: {filename}")
    time.sleep(2)

    text = extract_text(pdf_path)
    if not text:
        print(f"[!] No se pudo extraer texto: {filename}")
        return

    print("[*] Analizando criterios...")
    analysis = analyze_pcap(text)

    print("[*] Generando Word...")
    out_path = os.path.join(OUTPUT_DIR, f"{name_only}_Criterios.docx")
    create_word_report(name_only, analysis, output_path=out_path)
    print(f"[+] Informe: {out_path}")

    if analysis.get('warnings'):
        for w in analysis['warnings']:
            print(f"[!] {w}")

    try:
        shutil.move(pdf_path, os.path.join(PROCESSED_DIR, filename))
        print(f"[*] PDF movido a procesados/")
    except Exception as e:
        print(f"[!] Error moviendo: {e}")


try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler

    class PCAPHandler(FileSystemEventHandler):
        def on_created(self, event):
            if not event.is_directory and event.src_path.lower().endswith(".pdf"):
                process_single_pdf(event.src_path)

    def start_watching():
        setup_directories()
        for file in os.listdir(INPUT_DIR):
            if file.lower().endswith(".pdf"):
                process_single_pdf(os.path.join(INPUT_DIR, file))

        event_handler = PCAPHandler()
        observer = Observer()
        observer.schedule(event_handler, INPUT_DIR, recursive=False)
        observer.start()
        print(f"\n[*] Carpeta de escucha: {INPUT_DIR}")
        print("[*] Esperando PDFs... (Ctrl+C para detener)")
        try:
            while True:
                time.sleep(1)
        except KeyboardInterrupt:
            observer.stop()
        observer.join()

except ImportError:
    def start_watching():
        print("[!] watchdog no instalado. Usa la interfaz Streamlit.")


if __name__ == '__main__':
    start_watching()
